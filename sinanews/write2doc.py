# -*- coding: utf-8 -*-
import pymongo
from io import BytesIO
from PIL import Image
from docx.shared import Inches
from docx import Document
import requests

class Write2docx(object):
    def __init__(self):
        self.document = Document()
        self.mongo_collection = pymongo.MongoClient()['sinanews']['info']
    def _download_img(self, url):
        imgdata = requests.get(url).content
        return imgdata

    def _getdata(self):
        title = self.mongo_collection.find()
        return title


    def write_img(self,url):
        imgdata = self._download_img(url)
        image_io = BytesIO()
        image_io.write(imgdata)
        image_io.seek(0)
        im = Image.open(image_io)
        self.document.add_picture(image_io,width=Inches(3))

if __name__ == '__main__':
    test = Write2docx()
    for i in test._getdata():
        data = i['内容']
        for x in data:
            if x.startswith('http'):
                test.write_img(x)
            else:
                 p=test.document.add_paragraph(x)
    test.document.save('kkk.docx')
